import io
import logging

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

TYPE_LABELS = {
    'single_choice': '单选题',
    'multi_choice': '多选题',
    'fill_blank': '填空题',
    'true_false': '判断题',
    'calculation': '计算题',
    'short_answer': '简答题',
    'essay': '论述题',
}

DIFFICULTY_LABELS = {
    'easy': '简单',
    'medium': '中等',
    'hard': '困难',
}


def export_paper_to_docx_bytes(paper_data: dict) -> bytes:
    doc = Document()

    _setup_styles(doc)

    _add_title(doc, paper_data)

    _add_questions(doc, paper_data.get('questions', paper_data.get('paperJson', {}).get('questions', [])))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def export_paper_from_model(paper, questions) -> bytes:
    doc = Document()

    _setup_styles(doc)

    title = paper.paperTitle or paper.courseName or '试卷'
    doc.add_heading(str(title), level=0)

    total_score = sum(q.score or 0 for q in questions)
    meta_text = f'课程：{paper.courseName or "未命名"}    |    共 {len(questions)} 题    |    总分 {total_score} 分'
    p = doc.add_paragraph(meta_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        run = p.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph('')

    question_list = []
    for pq in questions:
        q = {
            'question_no': pq.questionNo or 0,
            'question_type': pq.questionType or '',
            'content': pq.content or '',
            'options': _normalize_options(pq.options),
            'answer': pq.answer or '',
            'analysis': pq.analysis or '',
            'knowledge_points': pq.knowledgePoints or [],
            'difficulty': pq.difficulty or 'medium',
            'score': pq.score or 0,
        }
        question_list.append(q)

    _add_questions(doc, question_list)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _setup_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    if rpr.find(qn('w:rFonts')) is None:
        from lxml import etree
        etree.SubElement(rpr, qn('w:rFonts'))
    rpr.find(qn('w:rFonts')).set(qn('w:eastAsia'), '宋体')


def _add_title(doc, paper_data):
    paper_json = paper_data.get('paperJson', paper_data)
    title = paper_json.get('paper_title', paper_data.get('paperTitle', ''))
    if not title:
        title = f"{paper_data.get('courseName', '未命名')} 复习试卷"

    doc.add_heading(str(title), level=0)

    meta_parts = []
    course = paper_data.get('courseName', paper_json.get('course_name', ''))
    if course:
        meta_parts.append(f'课程：{course}')
    questions = paper_json.get('questions', [])
    if questions:
        total_score = sum(q.get('score', 0) for q in questions)
        meta_parts.append(f'共 {len(questions)} 题')
        meta_parts.append(f'总分 {total_score} 分')

    if meta_parts:
        p = doc.add_paragraph('    |    '.join(meta_parts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph('')


def _add_questions(doc, questions):
    questions_by_type = {}
    for q in questions:
        qt = q.get('question_type', '')
        questions_by_type.setdefault(qt, []).append(q)

    overall_index = 1
    for qt, qlist in questions_by_type.items():
        type_label = TYPE_LABELS.get(qt, qt)
        type_total_score = sum(q.get('score', 0) for q in qlist)

        heading = doc.add_heading(f'{type_label}（共 {len(qlist)} 题，合计 {type_total_score} 分）', level=2)

        for q in qlist:
            _add_single_question(doc, q, overall_index)
            overall_index += 1


def _add_single_question(doc, q, index):
    question_no = q.get('question_no', index)
    difficulty = q.get('difficulty', 'medium')
    score = q.get('score', 0)
    content = q.get('content', '')
    qtype = q.get('question_type', '')

    header_text = f'第{question_no}题  [{DIFFICULTY_LABELS.get(difficulty, difficulty)}]  （{score}分）'
    p = doc.add_paragraph()
    run = p.add_run(header_text)
    run.bold = True
    run.font.size = Pt(11)

    content_para = doc.add_paragraph(content)
    if content_para.runs:
        content_para.runs[0].font.size = Pt(11)

    options = _normalize_options(q.get('options'))
    if options and qtype in ('single_choice', 'multi_choice'):
        for opt in options:
            opt_text = f'{opt.get("key", "")}. {opt.get("value", "")}'
            opt_para = doc.add_paragraph(opt_text)
            opt_para.paragraph_format.left_indent = Inches(0.3)
            if opt_para.runs:
                opt_para.runs[0].font.size = Pt(11)

    answer = q.get('answer', '')
    if answer:
        answer_para = doc.add_paragraph()
        answer_run = answer_para.add_run(f'【答案】{answer}')
        answer_run.bold = True
        answer_run.font.size = Pt(11)
        answer_run.font.color.rgb = RGBColor(0, 100, 0)

    analysis = q.get('analysis', '')
    if analysis:
        analysis_para = doc.add_paragraph()
        analysis_run = analysis_para.add_run(f'【解析】{analysis}')
        analysis_run.font.size = Pt(10)
        analysis_run.font.color.rgb = RGBColor(80, 80, 80)

    knowledge_points = q.get('knowledge_points', [])
    if knowledge_points:
        kp_text = '、'.join(knowledge_points)
        kp_para = doc.add_paragraph()
        kp_run = kp_para.add_run(f'【知识点】{kp_text}')
        kp_run.font.size = Pt(9)
        kp_run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph('')


def _normalize_options(options):
    if not options:
        return []
    if isinstance(options, list):
        return options
    if isinstance(options, dict):
        return [{'key': k, 'value': v} for k, v in options.items()]
    return []
